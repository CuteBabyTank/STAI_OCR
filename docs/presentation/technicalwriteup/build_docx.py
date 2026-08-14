"""Build Snag_Technical_Writeup.docx from TECHNICAL_WRITEUP.md.

The write-up is authored once, in Markdown. This renders that file into the
existing `Snag_Technical_Writeup.docx` -- opening the document as its own
style template, emptying the body, and writing the current text back into it,
so headings, fonts and page setup stay exactly what they already were.

The ASCII-art blocks in the Markdown do not survive the trip. Each one is
replaced by the drawn figure that says the same thing, and seven more figures
are inserted where the prose reaches a result worth seeing. Figures come from
`figures.py`; run that first, or let this script run it.

Usage:
    python docs/presentation/technicalwriteup/build_docx.py [--source PATH]
                                                            [--out PATH]
                                                            [--no-figures]

`--source` defaults to the Markdown sitting next to this file, which is the one
to edit. The .docx is output: hand-edits to it are lost on the next build.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
ROOT = HERE.parents[2]
ASSETS = HERE / "assets"

DEFAULT_SOURCE = HERE / "TECHNICAL_WRITEUP.md"
DEFAULT_OUT = ROOT / "Snag_Technical_Writeup.docx"

# The document's original styles, page setup and theme, kept as a pristine copy.
# Building from the *output* instead would work until Word re-saves it — Word
# renames the Google-Docs `normal` style to `Normal` and quadruples the file
# size, so the next build would either crash or inherit the bloat.
TEMPLATE = HERE / "_template.docx"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5F, 0x67, 0x72)
ACCENT = RGBColor(0x8A, 0x3F, 0x07)
RULE = "D9DEE4"
HEADER_FILL = "EEF1F4"
CODE_FILL = "F5F7F9"

CONTENT_WIDTH = Inches(6.0)          # page width less both margins
MAX_FIGURE_HEIGHT = Inches(7.6)      # leaves room for a caption on one page


# --------------------------------------------------------------------------
# Figures: what replaces an ASCII block, and what gets inserted beside prose.
#
# `after`/`before` keys are matched against the first characters of a block's
# own Markdown source, so an anchor breaks loudly (the figure goes missing)
# rather than silently landing in the wrong section.
# --------------------------------------------------------------------------

REPLACE_CODE_BLOCK = {
    "   Capture": (
        "fig-01-architecture.png",
        "Figure 1. Five blocks, and the three things that run under all of "
        "them. The API, the models and the traces are not a layer each block "
        "calls — they are the same surface every block is reached through.",
    ),
    "Receipt image/PDF": (
        "fig-02-pipeline.png",
        "Figure 2. The extraction pipeline. Shaded stages are guardrails; a "
        "receipt that fails the audit or the disambiguation check is held, not "
        "filed. Nothing on this path computes a figure the receipt did not "
        "print. The agent half of the flow is Figure 7.",
    ),
    "Question + the last 10 turns": (
        "fig-07-react-loop.png",
        "Figure 7. The ReAct loop. Generation is stopped at the token "
        "Observation:, so the model cannot write its own observation — that is "
        "what makes this a loop over tools rather than a monologue about them.",
    ),
}

INSERT_BEFORE = {
    "**Provenance, stated plainly.**": (
        "fig-03-three-models.png",
        "Figure 3. The three-model comparison, with its provenance drawn in "
        "rather than left to a footnote.",
    ),
    "Processing time is driven by": (
        "fig-04-per-receipt.png",
        "Figure 4. Time per receipt, taken from the benchmark's own "
        "`elapsed_s` values. These timings measure a 4 GB laptop card, not the "
        "model — quote them only alongside §7.3.",
    ),
    "**1. Phantom `subtotal`": (
        "fig-05-error-taxonomy.png",
        "Figure 5. Where the 26 errors are. A wrong value is charged as both a "
        "false positive and a false negative, which is why the two columns "
        "overlap on the identifier and VAT classes.",
    ),
    "### 8.1 The agent model was being evicted": (
        "fig-06-performance.png",
        "Figure 6. What the performance pass moved. One vision call is ~6,000 "
        "ms, so a removed round trip is worth roughly 400× more than any "
        "amount of Python — which is why the round-trip count is the thing "
        "with a regression test on it.",
    ),
    "Specifically on the write path:": (
        "fig-08-guardrails.png",
        "Figure 8. Seven gates between a request and the ledger. Each is a "
        "named function, and each drops a specific class of thing rather than "
        "scoring a request for general suspiciousness.",
    ),
    "Four design decisions worth calling out:": (
        "fig-09-deployment.png",
        "Figure 9. The deployment. Three containers, one command, and "
        "inference as an HTTP call out — which is why the whole stack comes up "
        "on a laptop with no GPU.",
    ),
    "**Two caveats stated rather than buried.**": (
        "fig-10-cost.png",
        "Figure 10. The same 300 receipts a month, three ways — and the human "
        "hours each way costs.",
    ),
}


# --------------------------------------------------------------------------
# Markdown -> blocks
# --------------------------------------------------------------------------

class Block:
    __slots__ = ("kind", "body", "meta", "src")

    def __init__(self, kind, body, meta=None, src=""):
        self.kind = kind          # heading | para | list | table | code | quote | rule
        self.body = body
        self.meta = meta or {}
        self.src = src

    def __repr__(self):           # pragma: no cover - debugging only
        return f"<{self.kind} {str(self.body)[:40]!r}>"


def parse_markdown(text: str) -> list[Block]:
    lines = text.split("\n")
    blocks: list[Block] = []
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # fenced code
        if line.lstrip().startswith("```"):
            lang = line.strip().strip("`").strip()
            i += 1
            buf = []
            while i < n and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            body = "\n".join(buf)
            blocks.append(Block("code", body, {"lang": lang}, src=body))
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line.strip()):
            blocks.append(Block("rule", None, src=line))
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            blocks.append(Block("heading", m.group(2).strip(),
                                {"level": len(m.group(1))}, src=line))
            i += 1
            continue

        # table
        if line.lstrip().startswith("|") and i + 1 < n and \
                re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            header = _split_row(line)
            aligns = _alignments(lines[i + 1])
            i += 2
            rows = []
            while i < n and lines[i].lstrip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            blocks.append(Block("table", (header, rows), {"aligns": aligns},
                                src=line))
            continue

        # blockquote
        if line.lstrip().startswith(">"):
            buf = []
            src = line
            while i < n and lines[i].lstrip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            blocks.append(Block("quote", _unwrap(buf), src=src))
            continue

        # list (bulleted or numbered), possibly with wrapped continuation lines
        m = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", line)
        if m:
            ordered = not m.group(2) in "-*+"
            items: list[tuple[int, str]] = []
            src = line
            while i < n:
                mm = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", lines[i])
                if mm:
                    items.append((len(mm.group(1)) // 2, mm.group(3).strip()))
                    i += 1
                elif lines[i].strip() and lines[i].startswith((" ", "\t")) and items:
                    lvl, txt = items[-1]
                    items[-1] = (lvl, txt + " " + lines[i].strip())
                    i += 1
                elif not lines[i].strip():
                    nxt = lines[i + 1] if i + 1 < n else ""
                    if re.match(r"^(\s*)([-*+]|\d+\.)\s+", nxt):
                        i += 1
                    else:
                        break
                else:
                    break
            blocks.append(Block("list", items, {"ordered": ordered}, src=src))
            continue

        # paragraph
        buf = []
        src = line
        while i < n and lines[i].strip() and not _starts_block(lines[i]):
            buf.append(lines[i].strip())
            i += 1
        blocks.append(Block("para", _unwrap(buf), src=src))

    return blocks


def _starts_block(line: str) -> bool:
    s = line.lstrip()
    return (s.startswith(("```", "|", ">", "#"))
            or bool(re.match(r"^([-*+]|\d+\.)\s+", s))
            or bool(re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line.strip())))


def _unwrap(buf: list[str]) -> str:
    return " ".join(x.strip() for x in buf).strip()


def _split_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    parts, cur, esc = [], "", False
    for ch in line:
        if esc:
            cur += ch
            esc = False
        elif ch == "\\":
            esc = True
        elif ch == "|":
            parts.append(cur.strip())
            cur = ""
        else:
            cur += ch
    parts.append(cur.strip())
    return parts


def _alignments(sep: str) -> list[str]:
    out = []
    for cell in _split_row(sep):
        left, right = cell.startswith(":"), cell.endswith(":")
        out.append("center" if left and right else "right" if right else "left")
    return out


# --------------------------------------------------------------------------
# Inline formatting
# --------------------------------------------------------------------------

INLINE = re.compile(
    r"(?P<code>`[^`]+`)"
    r"|(?P<bold>\*\*[^*]+\*\*)"
    r"|(?P<ital>(?<!\*)\*(?!\*)[^*]+\*(?!\*))"
    r"|(?P<sup><sup>.*?</sup>)"
    r"|(?P<link>\[[^\]]+\]\([^)]+\))"
)


def add_runs(paragraph, text: str, *, size=None, color=None, italic=False,
             bold=False):
    """Write Markdown inline formatting into a paragraph as real runs."""
    text = text.replace("<br>", " ").replace("&nbsp;", " ")
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], size=size, color=color,
                 italic=italic, bold=bold)
        kind = m.lastgroup
        raw = m.group()
        if kind == "code":
            _run(paragraph, raw[1:-1], size=size, color=ACCENT, mono=True,
                 bold=bold)
        elif kind == "bold":
            # recurse: `code` and *emphasis* nest inside **bold** in the source
            add_runs(paragraph, raw[2:-2], size=size, color=color, bold=True,
                     italic=italic)
        elif kind == "ital":
            add_runs(paragraph, raw[1:-1], size=size, color=color, italic=True,
                     bold=bold)
        elif kind == "sup":
            inner = re.sub(r"</?sup>", "", raw)
            r = _run(paragraph, inner, size=size, color=color)
            r.font.superscript = True
        elif kind == "link":
            label, target = re.match(r"\[([^\]]+)\]\(([^)]+)\)", raw).groups()
            if target.startswith("#"):          # an in-document cross-reference
                add_runs(paragraph, label, size=size, color=color,
                         italic=italic, bold=bold)
            else:
                _run(paragraph, label, size=size, color=color, italic=italic,
                     bold=bold)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], size=size, color=color, italic=italic,
             bold=bold)


def _run(paragraph, text, *, size=None, color=None, bold=False, italic=False,
         mono=False):
    r = paragraph.add_run(text)
    if size:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    if mono:
        r.font.name = "Consolas"
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(attr), "Consolas")
        if size is None:
            r.font.size = Pt(9.5)
    return r


# --------------------------------------------------------------------------
# Low-level docx helpers
# --------------------------------------------------------------------------

def _shade(element, fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _spacing(paragraph, *, before=None, after=None, line=None):
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _table_borders(table, colour=RULE, size=4):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), colour)
        borders.append(el)
    tbl_pr.append(borders)


def _keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    ppr.append(el)


def _repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    tr_pr.append(el)


# --------------------------------------------------------------------------
# Renderers
# --------------------------------------------------------------------------

class Writer:
    def __init__(self, doc):
        self.doc = doc
        self.figure_count = 0
        # Google Docs writes `normal`, Word writes `Normal`. Resolve by name,
        # case-insensitively, so either template builds.
        self._styles = {s.name.lower(): s.name for s in doc.styles if s.name}
        self.body_style = self._style("normal")

    def _style(self, wanted: str) -> str:
        name = self._styles.get(wanted.lower())
        if name is None:
            raise SystemExit(
                f"the template has no {wanted!r} style — has it been re-saved "
                f"by another editor? styles present: {sorted(self._styles)}")
        return name

    # -- text ------------------------------------------------------------
    def heading(self, text, level):
        style = self._style(
            {1: "Title", 2: "Heading 1", 3: "Heading 2"}.get(level, "Heading 3"))
        p = self.doc.add_paragraph(style=style)
        add_runs(p, text)
        if level > 1:
            _spacing(p, before=16 if level == 2 else 12, after=6)
        return p

    def para(self, text, *, size=None, color=None, italic=False, indent=0):
        p = self.doc.add_paragraph(style=self.body_style)
        add_runs(p, text, size=size, color=color, italic=italic)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        _spacing(p, after=8, line=1.28)
        return p

    def bullets(self, items, ordered=False):
        for level, text in items:
            p = self.doc.add_paragraph(style=self.body_style)
            p.paragraph_format.left_indent = Inches(0.28 + 0.28 * level)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            marker = "· " if not ordered else ""
            if marker:
                _run(p, marker, color=MUTED)
            add_runs(p, text)
            _spacing(p, after=3, line=1.22)

    def numbered(self, items):
        for idx, (level, text) in enumerate(items, start=1):
            p = self.doc.add_paragraph(style=self.body_style)
            p.paragraph_format.left_indent = Inches(0.34)
            p.paragraph_format.first_line_indent = Inches(-0.34)
            _run(p, f"{idx}.  ", color=MUTED)
            add_runs(p, text)
            _spacing(p, after=3, line=1.22)

    def quote(self, text):
        p = self.doc.add_paragraph(style=self.body_style)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.2)
        add_runs(p, text, size=9.5, color=MUTED)
        _spacing(p, before=4, after=10, line=1.25)
        ppr = p._p.get_or_add_pPr()
        bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "12")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "C2C9D2")
        bdr.append(left)
        ppr.append(bdr)

    def code(self, text):
        for line in text.split("\n"):
            p = self.doc.add_paragraph(style=self.body_style)
            _run(p, line or " ", size=8.5, color=INK, mono=True)
            _spacing(p, before=0, after=0, line=1.0)
            _shade(p._p.get_or_add_pPr(), CODE_FILL)
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.right_indent = Inches(0.12)
        trailer = self.doc.add_paragraph(style=self.body_style)
        _spacing(trailer, before=0, after=6, line=1.0)
        _run(trailer, "", size=4)

    def rule(self):
        p = self.doc.add_paragraph(style=self.body_style)
        _spacing(p, before=2, after=8)
        _run(p, "", size=2)
        ppr = p._p.get_or_add_pPr()
        bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), RULE)
        bdr.append(bottom)
        ppr.append(bdr)

    # -- table -----------------------------------------------------------
    def table(self, header, rows, aligns):
        cols = max(len(header), max((len(r) for r in rows), default=0))
        # `| | |` in the source is a key/value table, not a headed one — an
        # empty shaded strip across the top would only look like a mistake.
        headed = any(c.strip() for c in header)

        t = self.doc.add_table(rows=1 if headed else 0, cols=cols)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = True
        _table_borders(t)

        if headed:
            for j in range(cols):
                cell = t.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs(p, header[j] if j < len(header) else "", size=9.5)
                for r in p.runs:
                    r.bold = True
                _spacing(p, before=3, after=3, line=1.1)
                _shade(cell._tc.get_or_add_tcPr(), HEADER_FILL)
            _repeat_header(t.rows[0])

        for row in rows:
            cells = t.add_row().cells
            for j in range(cols):
                cell = cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs(p, row[j] if j < len(row) else "", size=9.5)
                _spacing(p, before=3, after=3, line=1.1)
                if j < len(aligns) and aligns[j] == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif j < len(aligns) and aligns[j] == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        after = self.doc.add_paragraph(style=self.body_style)
        _spacing(after, before=0, after=8, line=1.0)
        _run(after, "", size=4)

    # -- figure ----------------------------------------------------------
    def figure(self, filename, caption):
        path = ASSETS / filename
        if not path.exists():
            raise FileNotFoundError(
                f"{path} is missing — run figures.py first")
        from PIL import Image
        with Image.open(path) as im:
            w, h = im.size
        width = CONTENT_WIDTH
        if Inches(width.inches * h / w) > MAX_FIGURE_HEIGHT:
            width = Inches(MAX_FIGURE_HEIGHT.inches * w / h)

        p = self.doc.add_paragraph(style=self.body_style)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, before=10, after=4, line=1.0)
        p.add_run().add_picture(str(path), width=width)
        _keep_with_next(p)

        cap = self.doc.add_paragraph(style=self.body_style)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(cap, before=0, after=12, line=1.2)
        cap.paragraph_format.left_indent = Inches(0.35)
        cap.paragraph_format.right_indent = Inches(0.35)
        add_runs(cap, caption, size=8.5, color=MUTED, italic=True)
        self.figure_count += 1


# --------------------------------------------------------------------------
# Build
# --------------------------------------------------------------------------

def _empty_body(doc):
    """Keep the section properties, drop everything else."""
    body = doc.element.body
    for child in list(body.iterchildren()):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _drop_orphan_images(doc):
    """Old figures are no longer referenced; don't ship their bytes."""
    part = doc.part
    used = {el.get(qn("r:embed")) for el in
            doc.element.body.iter(qn("a:blip"))}
    used |= {el.get(qn("r:link")) for el in doc.element.body.iter(qn("a:blip"))}
    for rid, rel in list(part.rels.items()):
        if "image" in rel.reltype and rid not in used:
            part.drop_rel(rid)


def _cover(writer, blocks):
    """The title page keeps the document's own framing, not the file's H1."""
    writer.heading("Snag: Receipt Processor and Ledger Agent", 1)
    p = writer.para("Technical Write-Up", size=13, color=MUTED)
    _spacing(p, after=2)
    writer.para(
        "An AI app that reads receipts, tracks expenses, and lets you ask "
        "questions about your spending in plain English.",
        size=11, color=MUTED, italic=True)


def build(source: Path, out: Path, *, run_figures: bool = True) -> Path:
    if run_figures:
        sys.path.insert(0, str(HERE))
        import figures
        figures.build_all()

    text = source.read_text(encoding="utf-8")
    blocks = parse_markdown(text)

    base = TEMPLATE if TEMPLATE.exists() else (out if out.exists() else None)
    doc = docx.Document(str(base)) if base else docx.Document()
    _empty_body(doc)
    w = Writer(doc)

    _cover(w, blocks)

    used_replacements: set[str] = set()
    used_inserts: set[str] = set()
    skip_title = True

    for block in blocks:
        # a figure that belongs immediately before this block
        for anchor, (img, caption) in INSERT_BEFORE.items():
            if anchor in used_inserts:
                continue
            if block.src.strip().startswith(anchor):
                w.figure(img, caption)
                used_inserts.add(anchor)

        if block.kind == "heading":
            if block.meta["level"] == 1 and skip_title:
                skip_title = False          # the cover already carries it
                continue
            w.heading(block.body, block.meta["level"])

        elif block.kind == "para":
            w.para(block.body)

        elif block.kind == "list":
            if block.meta["ordered"]:
                w.numbered(block.body)
            else:
                w.bullets(block.body)

        elif block.kind == "table":
            header, rows = block.body
            w.table(header, rows, block.meta["aligns"])

        elif block.kind == "quote":
            w.quote(block.body)

        elif block.kind == "rule":
            w.rule()

        elif block.kind == "code":
            swapped = False
            for anchor, (img, caption) in REPLACE_CODE_BLOCK.items():
                if block.body.lstrip("\n").startswith(anchor):
                    w.figure(img, caption)
                    used_replacements.add(anchor)
                    swapped = True
                    break
            if not swapped:
                w.code(block.body)

    missing_r = set(REPLACE_CODE_BLOCK) - used_replacements
    missing_i = set(INSERT_BEFORE) - used_inserts
    if missing_r or missing_i:
        raise SystemExit(
            "Anchors did not match the source — the Markdown moved:\n"
            + "".join(f"  code block: {a!r}\n" for a in sorted(missing_r))
            + "".join(f"  insert before: {a!r}\n" for a in sorted(missing_i)))

    _drop_orphan_images(doc)
    doc.save(str(out))
    return out


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--source", type=Path, default=DEFAULT_SOURCE,
                    help="the Markdown write-up to render")
    ap.add_argument("--out", type=Path, default=DEFAULT_OUT,
                    help="the .docx to rewrite in place")
    ap.add_argument("--no-figures", action="store_true",
                    help="use the PNGs already in assets/")
    args = ap.parse_args()

    if not args.source.exists():
        raise SystemExit(f"no such source: {args.source}")

    out = build(args.source, args.out, run_figures=not args.no_figures)
    doc = docx.Document(str(out))
    figures = sum(1 for _ in doc.element.body.iter(qn("a:blip")))
    print(f"\nWrote {out}")
    print(f"  source     {args.source}")
    print(f"  paragraphs {len(doc.paragraphs)}")
    print(f"  tables     {len(doc.tables)}")
    print(f"  figures    {figures}")


if __name__ == "__main__":
    main()
